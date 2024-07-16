import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.shared import Pt,Cm,Inches
import re

#Data cleaning
def CJK_cleaner(string): 
    filters = re.compile('[^0-9a-zA-Z\u4e00-\u9fff\uFF00-\uFFEF\u3000-\u303F\uFF01-\uFF0F\uFF1A-\uFF20\uFF3B-\uFF40\uFF5B-\uFF65]+', re.UNICODE)

    return filters.sub('', string) #remove special characters

#New Document
def GenerateNewWord(filename):
    document = Document()
    document.save(filename)

#Document setting
def document_set(document):
    document.styles['Normal'].font.name = u'新細明體'
    #獲取邊界
    #sec0=document.sections[0]
    #print("left margin:",sec0.left_margin.cm)
    #print("right margin:",sec0.right_margin.cm)
    #print("width :",sec0.page_width.cm)
    
#標題style
def document_style(document):
    style=document.styles.add_style('textstyle', WD_STYLE_TYPE. PARAGRAPH)
    style.font.size=Pt(22) #大小
    style.font.bold =True #粗體
    style.font.name = u'新細明體'

    return style
    
#main
while(True):

  #example input url='https://udn.com/news/story/6813/7066523?from=udn-referralnews_ch2artbottom'
  url1=""
  url1=input("please input")

  #GET
  try:
   html = requests.get(url1)
  except:
   print("FAILED")
   continue

  #UTF-8 中文
  html.encoding = 'UTF-8'

  #BS4
  sp = BeautifulSoup(html.text, 'html5lib')

  #打開文件
  try:
   document = Document('聯合新聞網TEST.docx')
  except:
   print("FAILED GET FILE")
   continue

  #Document setting  ########################################
  document_set(document)
  style=document_style(document)

  #Document content  ########################################
  #新增title
  try:
   title=sp.find('h1',class_="article-content__title").text.replace(" "," ")
   print(title)
   p = document.add_paragraph(text=title,style=style)
  # p.paragraph_format.hyphenation = True 
  except:
   title=sp.find('h1').text.replace(" "," ")
   print(title)
   p = document.add_paragraph(text=title,style=style)
   p.paragraph_format.hyphenation = True

  #放入作者
  try:
    au=""
    for i in sp.find('section',class_='authors'):
      au=au+i.text
    au=au.replace("\n","")  
    au=au.replace(" ","")
    print("作者:",au)
    p = document.add_paragraph(au)
  except:
    print("authors error")

  #放入內容
  count = 0
  for i in sp.select('section.article-content__editor > p'):# sp.find('section',class_="article-content__editor").find_all('p'):
     if i.text == "\n" or len(i.text)==0:
      count = count+ 1
     else:
      if count==0:
        print("X:line")
        #data add space
        p = document.add_paragraph("") 
      count=0

     #Except Data cleaning 
     if count>1:
        continue
     print(len(i.text),"  ","X:",i.text.replace("\n",""))
     if(i.text.find("【延伸閱讀】")!=-1):
        break
     if(i.text.find("※ 提醒您：禁止酒駕 飲酒過量有礙健康")!=-1):
        continue
     if(i.text.find("★珍惜生命，若您或身邊的人有心理困擾，可撥打")!=-1):
        continue
     if(i.text.find("© ")!=-1):
        continue
     
     #data add
     p = document.add_paragraph(CJK_cleaner(i.text))#.replace("\n","").replace("\t","").replace("\r",""))
  
  #儲存名字
  document.save(f'news/{title}.doc')
  #Document content  ########################################
